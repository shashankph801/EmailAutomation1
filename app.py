import os
import re
import html
from functools import wraps
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email.mime.application import MIMEApplication
from email import encoders
from flask import Flask, request, jsonify, render_template, redirect, url_for, session
from werkzeug.utils import secure_filename
import pypdf
from docx import Document
import datetime
from dotenv import load_dotenv

import uuid

from google_auth_oauthlib.flow import Flow
from gmail_client import credentials_ready, gmail_api_configured, send_via_gmail_api
from db import get_db, get_hr_user

# load local .env into environment (optional; safe, .env is in .gitignore)
load_dotenv(dotenv_path=".env", override=True)

if os.environ.get('SENDER_EMAIL'):
    os.environ['SENDER_EMAIL'] = os.environ['SENDER_EMAIL'].strip()

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 20MB max
app.secret_key = os.environ.get('SECRET_KEY', 'change-me-in-production')

from werkzeug.middleware.proxy_fix import ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'email' not in session:
            if request.is_json:
                return jsonify({'success': False, 'error': 'Not logged in'}), 401
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

UPLOAD_FOLDER = 'uploads'
ATTACH_FOLDER = 'attachments'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(ATTACH_FOLDER, exist_ok=True)

ALLOWED_RESUME = {'pdf', 'docx'}
ALLOWED_ATTACH = {'pdf', 'docx', 'xlsx', 'xls', 'csv', 'txt', 'png', 'jpg', 'jpeg'}

EMAIL_REGEX = re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}')

DEFAULT_SETTINGS = {
    'roles': [
        {
            'id': 'ba',
            'name': 'BA',
            'upload_link': 'https://drive.google.com/drive/folders/1adqT4v0-L3A8FTMtUg2AbQCzRVpJ4Iy8',
            'email_subject': 'Internship Opportunity - Case Study Submission',
            'email_body': (
                'Dear Candidate,\n\n'
                'Thank you for your interest in the Internship Opportunity with us.\n\n'
                'As part of the selection process, please complete the attached case study and submit this within 3 business days.\n\n'
                'JD and Assignment has been attached in the drive link provided.\n\n'
                'This is a great opportunity for you to showcase your analytical thinking, creativity and problem-solving approach. '
                'We encourage you to put forward your best work.\n\n'
                'Further details regarding the process will be shared later.\n\n'
                'We look forward to reviewing your submission and wish you all the very best!\n\n'
                'Thank you,\n\nRegards\nHR Team'
            ),
            'invite_subject': 'Invitation: Business Analyst Internship Interview',
            'invite_body': (
                'Dear Candidate,\n\n'
                'Thank you for your interest in the Business Analyst Internship opportunity with PriceEasy.\n\n'
                'We are pleased to inform you that you have been shortlisted for the interview process.\n\n'
                'Please find the interview details.'
            )
        },
        {
            'id': 'dev',
            'name': 'Dev',
            'upload_link': 'https://drive.google.com/drive/folders/1qhQ8YS_hA1ZPbxIG8DqMxHfrtx2WeEKk',
            'email_subject': 'Application: Senior Frontend Developer',
            'email_body': (
                'Dear Candidate,\n\n'
                'Thank you for your interest in the Internship Opportunity with us.\n\n'
                'As part of the selection process, please complete the attached case study and submit this within 3 business days.\n\n'
                'JD and Assignment has been attached in the drive link provided.\n\n'
                'This is a great opportunity for you to showcase your analytical thinking, creativity and problem-solving approach. '
                'We encourage you to put forward your best work.\n\n'
                'Further details regarding the process will be shared later.\n\n'
                'We look forward to reviewing your submission and wish you all the very best!\n\n'
                'Thank you,\n\nRegards\nHR Team'
            ),
            'invite_subject': 'Invitation: Software Developer Interview',
            'invite_body': (
                'Dear Candidate,\n\n'
                'Thank you for your interest in software development with PriceEasy.\n\n'
                'We are pleased to inform you that you have been shortlisted for the interview process.\n\n'
                'Please find the interview details.'
            )
        }
    ],
    'panels': [
        {'id': 'shashank', 'name': 'Shashank', 'email': 'shashankphshetty@gmail.com'}
    ]
}


def allowed_file(filename, allowed):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed


def extract_text_from_pdf(filepath):
    text = ''
    try:
        reader = pypdf.PdfReader(filepath)
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + '\n'
    except Exception as e:
        raise RuntimeError(f'PDF read error: {e}')
    return text


def extract_text_from_docx(filepath):
    text = ''
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            text += para.text + '\n'
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + ' '
                text += '\n'
    except Exception as e:
        raise RuntimeError(f'DOCX read error: {e}')
    return text


def extract_email(text):
    matches = EMAIL_REGEX.findall(text)
    # Filter out common false positives
    blacklist = {'example.com', 'test.com', 'email.com', 'domain.com', 'yourname.com'}
    for m in matches:
        domain = m.split('@')[1].lower()
        if domain not in blacklist:
            return m
    return None


def slugify(value):
    slug = re.sub(r'[^a-z0-9]+', '-', value.lower()).strip('-')
    return slug or str(uuid.uuid4())


def seed_defaults_if_empty():
    db = get_db()
    if db.roles.count_documents({}) == 0:
        db.roles.insert_many([dict(r) for r in DEFAULT_SETTINGS['roles']])
    if db.panels.count_documents({}) == 0:
        db.panels.insert_many([dict(p) for p in DEFAULT_SETTINGS['panels']])


def read_settings():
    db = get_db()
    roles = list(db.roles.find({}, {'_id': 0}))
    panels = list(db.panels.find({}, {'_id': 0}))
    return {'roles': roles, 'panels': panels}


seed_defaults_if_empty()


def send_email(sender_email, to_email, subject, body, attachments, refresh_token=None):
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = to_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    for filepath, filename in attachments:
        if not os.path.exists(filepath):
            continue
        with open(filepath, 'rb') as f:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
        msg.attach(part)

    send_via_gmail_api(msg, refresh_token)


def _ics_escape(value):
    if not value:
        return ''
    return (
        value.replace('\\', '\\\\')
        .replace(';', '\\;')
        .replace(',', '\\,')
        .replace('\r\n', '\\n')
        .replace('\n', '\\n')
    )


def send_calendar_invite(sender_email, organiser_email, attendees, subject, body_text,
                         dtstart_utc, dtend_utc, location, attachments=None,
                         meet_link=None, phone=None, pin=None, refresh_token=None):
    # Gmail requires the organizer to match the authenticated sending account.
    organiser_email = sender_email

    msg = MIMEMultipart('mixed')
    msg['From'] = sender_email
    msg['To'] = ','.join(attendees)
    msg['Subject'] = subject
    # ICS DESCRIPTION: body text only. Do NOT put meet link here —
    # Gmail detects meet.google.com URLs in DESCRIPTION and renders a second Meet widget.
    # Meet link goes in LOCATION only, which calendar apps display without creating a widget.
    ics_description = _ics_escape((body_text or '').strip())
    event_location = location or meet_link or ''

    # Prepare UID and timestamps
    uid = f"{uuid.uuid4()}@{sender_email.split('@')[-1]}"
    dtstamp = datetime.datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')
    dtstart = dtstart_utc.strftime('%Y%m%dT%H%M%SZ')
    dtend = dtend_utc.strftime('%Y%m%dT%H%M%SZ')

    ics = (
        f"BEGIN:VCALENDAR\r\n"
        f"PRODID:-//ResumeAgent//EN\r\n"
        f"VERSION:2.0\r\n"
        f"CALSCALE:GREGORIAN\r\n"
        f"METHOD:REQUEST\r\n"
        f"BEGIN:VEVENT\r\n"
        f"UID:{uid}\r\n"
        f"DTSTAMP:{dtstamp}\r\n"
        f"DTSTART:{dtstart}\r\n"
        f"DTEND:{dtend}\r\n"
        f"SUMMARY:{_ics_escape(subject)}\r\n"
        f"DESCRIPTION:{ics_description}\r\n"
        f"LOCATION:{_ics_escape(event_location)}\r\n"
        f"ORGANIZER;CN={organiser_email}:mailto:{organiser_email}\r\n"
        f"TRANSP:OPAQUE\r\n"
        f"SEQUENCE:0\r\n"
        f"STATUS:CONFIRMED\r\n"
    )

    for a in attendees:
        ics += (
            f"ATTENDEE;CUTYPE=INDIVIDUAL;ROLE=REQ-PARTICIPANT;PARTSTAT=NEEDS-ACTION;"
            f"RSVP=TRUE;CN={a}:mailto:{a}\r\n"
        )

    ics += "END:VEVENT\r\nEND:VCALENDAR"

    # EMAIL 1: a PURE text/calendar message (no multipart, no attachments).
    # This is the only structure Gmail reliably renders as exactly ONE invite card.
    # The body text and meet link live in the ICS DESCRIPTION/LOCATION and show inside the card.
    # (Mixing the calendar with HTML body and a PDF attachment in one email makes Gmail
    #  render zero, one, or two invite cards unpredictably — so we keep it isolated.)
    invite_msg = MIMEText(ics, 'calendar', 'utf-8')
    invite_msg.replace_header('Content-Type', 'text/calendar; method=REQUEST; charset="UTF-8"')
    invite_msg['From'] = sender_email
    invite_msg['To'] = ','.join(attendees)
    invite_msg['Subject'] = subject

    invite_result = send_via_gmail_api(invite_msg, refresh_token)

    # EMAIL 2: optional follow-up with the resume / attachments as a normal email.
    attached_files = []
    attachment_result = None
    if attachments and any(p and os.path.exists(p) for p in attachments):
        att_msg = MIMEMultipart('mixed')
        att_msg['From'] = sender_email
        att_msg['To'] = ','.join(attendees)
        att_msg['Subject'] = f'Attachment for: {subject}'

        note = (body_text or '').strip()
        note = (note + '\n\n' if note else '') + 'Please find the attached document for the interview above.'
        att_msg.attach(MIMEText(note, 'plain', 'utf-8'))

        for path in attachments:
            if path and os.path.exists(path):
                fname = os.path.basename(path)
                attached_files.append({'path': path, 'name': fname, 'size': os.path.getsize(path)})
                with open(path, 'rb') as f:
                    data = f.read()
                subtype = 'pdf' if fname.lower().endswith('.pdf') else 'octet-stream'
                part = MIMEBase('application', subtype)
                part.set_payload(data)
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f'attachment; filename="{fname}"')
                att_msg.attach(part)

        if attached_files:
            attachment_result = send_via_gmail_api(att_msg, refresh_token)

    # EMAIL 3: ICS file attached to HR (sender) only so they can add the event to their own calendar.
    hr_result = {}
    try:
        hr_msg = MIMEMultipart('mixed')
        hr_msg['From'] = sender_email
        hr_msg['To'] = sender_email
        hr_msg['Subject'] = f'Calendar file: {subject}'
        hr_msg.attach(MIMEText(
            'Click the attached .ics file to add this interview to your Google Calendar.',
            'plain', 'utf-8'
        ))
        ics_part = MIMEApplication(ics.encode('utf-8'), Name='invite.ics')
        ics_part['Content-Disposition'] = 'attachment; filename="invite.ics"'
        hr_msg.attach(ics_part)
        hr_result = send_via_gmail_api(hr_msg, refresh_token)
    except Exception as e:
        print(f'[EMAIL 3] Failed to send ICS to HR ({sender_email}): {e}')

    return {
        'gmail_message_id': invite_result.get('id'),
        'attachment_message_id': (attachment_result or {}).get('id'),
        'hr_calendar_message_id': hr_result.get('id'),
        'attached_files': attached_files,
    }


def _make_oauth_flow(state=None):
    redirect_uri = os.environ.get('REDIRECT_URI', 'http://localhost:5000/auth/callback')
    if redirect_uri.startswith('http://'):
        os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
    config = {
        'web': {
            'client_id': os.environ.get('GOOGLE_CLIENT_ID'),
            'client_secret': os.environ.get('GOOGLE_CLIENT_SECRET'),
            'auth_uri': 'https://accounts.google.com/o/oauth2/auth',
            'token_uri': 'https://oauth2.googleapis.com/token',
            'redirect_uris': [redirect_uri],
        }
    }
    kwargs = {'scopes': ['https://www.googleapis.com/auth/gmail.send'], 'redirect_uri': redirect_uri}
    if state:
        kwargs['state'] = state
    return Flow.from_client_config(config, **kwargs)


@app.route('/login', methods=['GET', 'POST'])
def login():
    session.clear()
    error = None
    if request.method == 'POST':
        email = (request.form.get('email') or '').strip().lower()
        if not email:
            error = 'Email is required'
        else:
            hr_user = get_hr_user(email)
            if hr_user and hr_user.get('refresh_token'):
                session['email'] = email
                return redirect(url_for('main_page'))
            else:
                session['pending_email'] = email
                return redirect(url_for('setup_token_page'))
    return render_template('login.html', error=error)


@app.route('/setup-token')
def setup_token_page():
    email = session.get('pending_email', '')
    if not email:
        return redirect(url_for('login'))
    error = request.args.get('error', '')
    return render_template('setup_token.html', email=email, error=error)


@app.route('/auth/google')
def google_auth():
    email = session.get('pending_email')
    if not email:
        return redirect(url_for('login'))
    flow = _make_oauth_flow()
    authorization_url, state = flow.authorization_url(
        access_type='offline',
        prompt='consent',
        login_hint=email
    )
    session['oauth_state'] = state
    # Save the PKCE code_verifier so the callback can use the same one
    try:
        session['code_verifier'] = flow.code_verifier
    except Exception:
        session.pop('code_verifier', None)
    return redirect(authorization_url)


@app.route('/auth/callback')
def google_callback():
    state = session.get('oauth_state')
    email = session.get('pending_email')
    if not state or not email:
        return redirect(url_for('login'))
    try:
        flow = _make_oauth_flow(state=state)
        auth_response = request.url
        if request.headers.get('X-Forwarded-Proto') == 'https':
            auth_response = auth_response.replace('http://', 'https://', 1)
        # Pass the saved code_verifier so PKCE token exchange succeeds
        code_verifier = session.get('code_verifier')
        if code_verifier:
            flow.fetch_token(authorization_response=auth_response, code_verifier=code_verifier)
        else:
            flow.fetch_token(authorization_response=auth_response)
    except Exception as e:
        return redirect(url_for('setup_token_page') + '?error=' + html.escape(str(e)))
    refresh_token = flow.credentials.refresh_token
    if not refresh_token:
        return redirect(url_for('setup_token_page') + '?error=No+refresh+token+received.+Revoke+app+access+at+myaccount.google.com/permissions+and+try+again.')
    db = get_db()
    db.hr_users.update_one(
        {'email': email},
        {'$set': {'email': email, 'refresh_token': refresh_token}},
        upsert=True
    )
    session['email'] = email
    session.pop('pending_email', None)
    session.pop('oauth_state', None)
    return redirect(url_for('main_page'))


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


@app.route('/')
def index():
    return redirect(url_for('login'))


@app.route('/main')
@login_required
def main_page():
    return render_template('main.html')


@app.route('/app')
@login_required
def app_page():
    return render_template('index.html')


@app.route('/interview')
@login_required
def interview_page():
    return render_template('interview.html')


@app.route('/settings')
@login_required
def settings_page():
    return render_template('settings.html')


@app.route('/api/settings', methods=['GET'])
@login_required
def api_get_settings():
    return jsonify(read_settings())


@app.route('/api/roles', methods=['POST'])
@login_required
def api_add_role():
    data = request.json or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Role name is required'}), 400

    db = get_db()
    role_id = slugify(name)
    if db.roles.find_one(
        {'$or': [{'id': role_id}, {'name': re.compile(f'^{re.escape(name)}$', re.IGNORECASE)}]},
        {'_id': 0}
    ):
        return jsonify({'success': False, 'error': 'Role already exists'}), 409

    new_role = {
        'id': role_id,
        'name': name,
        'upload_link': (data.get('upload_link') or '').strip(),
        'email_subject': (data.get('email_subject') or f'Internship Opportunity - {name}').strip(),
        'email_body': (data.get('email_body') or DEFAULT_SETTINGS['roles'][0]['email_body']).strip(),
        'invite_subject': (data.get('invite_subject') or f'Invitation: {name} Interview').strip(),
        'invite_body': (data.get('invite_body') or DEFAULT_SETTINGS['roles'][0]['invite_body']).strip()
    }
    db.roles.insert_one(new_role)
    roles = list(db.roles.find({}, {'_id': 0}))
    return jsonify({'success': True, 'roles': roles}), 201


@app.route('/api/roles/<role_id>', methods=['DELETE'])
@login_required
def api_delete_role(role_id):
    db = get_db()
    result = db.roles.delete_one({'id': role_id})
    if result.deleted_count == 0:
        return jsonify({'success': False, 'error': 'Role not found'}), 404
    roles = list(db.roles.find({}, {'_id': 0}))
    return jsonify({'success': True, 'roles': roles})


@app.route('/api/panels', methods=['POST'])
@login_required
def api_add_panel():
    data = request.json or {}
    name = (data.get('name') or '').strip()
    email = (data.get('email') or '').strip()
    if not name or not email:
        return jsonify({'success': False, 'error': 'Panel name and email are required'}), 400
    if not EMAIL_REGEX.fullmatch(email):
        return jsonify({'success': False, 'error': 'Enter a valid email address'}), 400

    db = get_db()
    panel_id = slugify(name)
    if db.panels.find_one(
        {'email': re.compile(f'^{re.escape(email)}$', re.IGNORECASE)},
        {'_id': 0}
    ):
        return jsonify({'success': False, 'error': 'Panel email already exists'}), 409

    new_panel = {'id': panel_id, 'name': name, 'email': email}
    db.panels.insert_one(new_panel)
    panels = list(db.panels.find({}, {'_id': 0}))
    return jsonify({'success': True, 'panels': panels}), 201


@app.route('/api/panels/<panel_id>', methods=['DELETE'])
@login_required
def api_delete_panel(panel_id):
    db = get_db()
    result = db.panels.delete_one({'id': panel_id})
    if result.deleted_count == 0:
        return jsonify({'success': False, 'error': 'Panel not found'}), 404
    panels = list(db.panels.find({}, {'_id': 0}))
    return jsonify({'success': True, 'panels': panels})


@app.route('/extract-email', methods=['POST'])
@login_required
def extract_email_route():
    if 'resume' not in request.files:
        return jsonify({'success': False, 'error': 'No resume file provided'}), 400

    file = request.files['resume']
    if not file.filename or not allowed_file(file.filename, ALLOWED_RESUME):
        return jsonify({'success': False, 'error': 'Only PDF and DOCX resumes are supported'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    try:
        ext = filename.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            text = extract_text_from_pdf(filepath)
        else:
            text = extract_text_from_docx(filepath)

        if not text.strip():
            return jsonify({'success': False, 'error': 'Could not extract text from resume (may be image-based PDF)'}), 400

        email = extract_email(text)
        if not email:
            return jsonify({'success': False, 'error': 'No email address found in the resume'}), 404

        return jsonify({'success': True, 'email': email, 'resume_path': filepath})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/upload-attachment', methods=['POST'])
@login_required
def upload_attachment():
    key = request.form.get('key', 'file')
    if key not in request.files:
        return jsonify({'success': False, 'error': 'No file'}), 400
    file = request.files[key]
    if not file.filename:
        return jsonify({'success': False, 'error': 'Empty filename'}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(ATTACH_FOLDER, filename)
    file.save(filepath)
    return jsonify({'success': True, 'path': filepath, 'name': filename})


@app.route('/send', methods=['POST'])
@login_required
def send_route():
    data = request.json

    required = ['to_email', 'subject', 'body']
    for field in required:
        if not data.get(field):
            return jsonify({'success': False, 'error': f'Missing field: {field}'}), 400

    sender_email = session['email']
    hr_user = get_hr_user(sender_email)
    refresh_token = (hr_user or {}).get('refresh_token')
    ready, cred_error = credentials_ready(sender_email, refresh_token)
    if not ready:
        return jsonify({'success': False, 'error': cred_error}), 400

    attachments = []
    for att in data.get('attachments', []):
        path = att.get('path')
        name = att.get('name')
        if path and name and os.path.exists(path):
            attachments.append((path, name))

    try:
        send_email(
            sender_email=sender_email,
            to_email=data['to_email'],
            subject=data['subject'],
            body=data['body'],
            attachments=attachments,
            refresh_token=refresh_token
        )
        return jsonify({'success': True, 'message': f'Email sent to {data["to_email"]}'})
    except Exception as e:
        err = str(e)
        if 'invalid_grant' in err.lower():
            return jsonify({'success': False, 'error': 'Gmail API auth failed. Re-login or contact admin.'}), 401
        return jsonify({'success': False, 'error': err}), 500


@app.route('/send-invite', methods=['POST'])
@login_required
def send_invite_route():
    data = request.json or {}
    required = ['attendees', 'subject', 'invitation_text']
    for f in required:
        if not data.get(f):
            return jsonify({'success': False, 'error': f'Missing field: {f}'}), 400

    sender_email = session['email']
    hr_user = get_hr_user(sender_email)
    refresh_token = (hr_user or {}).get('refresh_token')
    ready, cred_error = credentials_ready(sender_email, refresh_token)
    if not ready:
        return jsonify({'success': False, 'error': cred_error}), 400

    organiser = sender_email
    attendees = [a.strip() for a in data['attendees'].split(',') if a.strip()]
    subject = data['subject']
    invitation_text = data['invitation_text']
    location = data.get('location', '')
    attachments = []
    # optional attachments field: single path or list
    if data.get('attachments'):
        if isinstance(data.get('attachments'), list):
            attachments = data.get('attachments')
        else:
            attachments = [data.get('attachments')]

    # parse date and times (date YYYY-MM-DD, start_time HH:MM, all in IST)
    # If date or times are missing/invalid, default to now+30min (IST) and 30min duration
    try:
        if data.get('date') and data.get('start_time') and data.get('end_time'):
            y,mm,dd = map(int, data['date'].split('-'))
            sh,sm = map(int, data['start_time'].split(':'))
            eh,em = map(int, data['end_time'].split(':'))
            # Build UTC datetimes from IST (IST = UTC +5:30)
            dtstart_utc = datetime.datetime(y, mm, dd, sh, sm) - datetime.timedelta(hours=5, minutes=30)
            dtend_utc = datetime.datetime(y, mm, dd, eh, em) - datetime.timedelta(hours=5, minutes=30)
            if dtend_utc <= dtstart_utc:
                dtend_utc = dtstart_utc + datetime.timedelta(minutes=30)
        else:
            # default: start = now +30min (UTC), end = start +30min (UTC)
            now_utc = datetime.datetime.utcnow()
            dtstart_utc = now_utc + datetime.timedelta(minutes=30)
            dtend_utc = dtstart_utc + datetime.timedelta(minutes=30)
    except Exception:
        # fallback to now+30min UTC
        now_utc = datetime.datetime.utcnow()
        dtstart_utc = now_utc + datetime.timedelta(minutes=30)
        dtend_utc = dtstart_utc + datetime.timedelta(minutes=30)

    try:
        invite_result = send_calendar_invite(
            sender_email=sender_email,
            organiser_email=organiser,
            attendees=attendees,
            subject=subject,
            body_text=invitation_text,
            dtstart_utc=dtstart_utc,
            dtend_utc=dtend_utc,
            location=location,
            attachments=attachments,
            meet_link=data.get('meet_link'),
            phone=data.get('phone'),
            pin=data.get('pin'),
            refresh_token=refresh_token
        )
        return jsonify({
            'success': True,
            'message': f'Invite sent to {len(attendees)} attendees',
        })
    except Exception as e:
        err = str(e)
        if 'invalid_grant' in err.lower():
            return jsonify({'success': False, 'error': 'Gmail API auth failed. Regenerate GOOGLE_REFRESH_TOKEN.'}), 401
        return jsonify({'success': False, 'error': err}), 500


@app.route('/env', methods=['GET'])
@login_required
def env_info():
    sender_email = session['email']
    hr_user = get_hr_user(sender_email)
    refresh_token = (hr_user or {}).get('refresh_token')
    ready, _ = credentials_ready(sender_email, refresh_token)
    return jsonify({
        'sender_email': sender_email,
        'credentials_ready': ready,
    })


if __name__ == '__main__':
    app.run(debug=False, port=5000)
